"""Microsoft Word (.docx) text extraction."""

from pathlib import Path

from docx import Document

from ingestion.models import ParsedDocument


def load_docx(path: Path) -> list[ParsedDocument]:
    """Return the document body and tables as individual extractable units."""
    document = Document(path)
    extracted: list[ParsedDocument] = []

    body = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    if body:
        extracted.append(
            ParsedDocument(body, path.name, "docx", "document body", {"section": "body"})
        )

    for table_number, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            extracted.append(
                ParsedDocument(
                    text="\n".join(rows),
                    source_name=path.name,
                    source_type="docx",
                    location=f"table {table_number}",
                    metadata={"table_number": table_number},
                )
            )
    return extracted

